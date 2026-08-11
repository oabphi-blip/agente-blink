"""Gera Word (.docx) com hyperlinks clicáveis — Consultas 16/06/2026."""
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "relatorio_consultas_16_06_2026.docx"

# ============================================================
# Helper hyperlinks clicáveis
# ============================================================

def add_hyperlink(paragraph, text, url, color="0000EE", underline=True):
    """Adiciona hyperlink clicável em um parágrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")  # 8pt
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rPr.append(rFonts)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_font(cell, size_pt=9, bold=False, color=None, align=None):
    for paragraph in cell.paragraphs:
        if align:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
            run.font.name = "Calibri"
            if bold:
                run.bold = True
            if color:
                run.font.color.rgb = color


# ============================================================
# Dados
# ============================================================

LINHAS = [
    ("08:30", "Nicolas Morales Monteiro", "2a", "Saúde Caixa", "Realizado", "8-Realizado", "24017039"),
    ("09:00", "Namera Roberta Souto Ribeiro", "37a", "TJDFT", "Realizado", "8-Realizado", "22604918"),
    ("09:30", "Giulliana Santos da Silva", "23a", "TRF", "Realizado", "8-Realizado", "22692648"),
    ("10:00", "Summaia Garzedin Santos de Abreu", "66a", "Serpro", "Realizado", "8-Realizado", "23863262"),
    ("14:00", "Thiago Andre Pierobom de Avila", "49a", "Plan Assiste", "Realizado", "8-Realizado", "22968530"),
    ("14:30", "Melissa Quintino Miranda ⚭", "7a", "Care Plus", "Realizado", "8-Realizado", "24141776"),
    ("15:00", "Maria Flor Quintino Miranda ⚭", "9a", "Care Plus", "Realizado", "8-Realizado", "24141776"),
    ("15:30", "Ryan Lucas Assunção Alves ⚭", "8a", "Saúde Caixa", "Realizado", "5-Agendado ⚠", "24153298"),
    ("16:00", "Iara Keile Assunção Silva Alves ⚭", "38a", "Saúde Caixa", "Realizado", "5-Agendado ⚠", "24153298"),
    ("16:00", "Mariana de Andrade Lima", "40a", "TJDFT", "Realizado", "(106157327)", "23943914"),
    ("16:30", "Gabriela Castro Silva", "18a", "Saúde Caixa", "Em atendimento", "Closed-won", "21768459"),
    ("17:00", "Bento Caetano dos Reis ⚠", "5 meses", "CORTESIA", "Agendado", "5-Agendado", "24118612"),
    ("17:00", "Milza de Castro Santana", "—", "Saúde Caixa", "Confirmado", "❌ Sem Kommo", ""),
    ("17:30", "Dominicky Ferreira Lemos", "9a", "Particular R$ 670", "Confirmado", "5-Agendado", "22345722"),
]

RESUMO = [
    ("Total agendamentos", "14"),
    ("Realizados", "10 (71%)"),
    ("Em atendimento", "1"),
    ("Confirmados aguardando chegar", "2"),
    ("Agendados sem confirmação", "1"),
    ("Cancelamentos", "0"),
]

ALERTAS = [
    ("1. Ryan + Iara (família) — etapa errada no Kommo",
     "Consultas realizadas no Medware (status 5), mas Kommo ainda em 5-AGENDADO. Mover pra 8-REALIZADO."),
    ("2. Milza de Castro Santana — sem lead Kommo",
     "Cadastrada direto no Medware, sem passagem pelo funil Kommo. Sem rastreabilidade."),
    ("3. Bento Caetano (5 meses) — anomalia de procedimento",
     "Procedimento código 309 = 'Paciente com 3 anos OU MAIS'. Marcado CORTESIA. Conferir cadastro."),
    ("4. Gabriela Castro Silva — 3 leads duplicados no Kommo",
     "Leads 11569512 (Closed-lost), 21768459 (Closed-won — ativo), 15036913 (Closed-lost). Sugere dedup."),
    ("5. 2 confirmados ainda não chegaram (17:00 Milza · 17:30 Dominicky)",
     "Recomenda ligação de confirmação de comparecimento."),
]


# ============================================================
# Documento
# ============================================================

def gerar():
    doc = Document()

    # Orientação paisagem A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    # ----- Cabeçalho -----
    titulo = doc.add_paragraph()
    titulo_run = titulo.add_run("Relação Completa — Consultas 16/06/2026")
    titulo_run.bold = True
    titulo_run.font.size = Pt(15)
    titulo_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Dra. Karla Delalíbera Pacheco · Águas Claras")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    obs = doc.add_paragraph()
    obs_run = obs.add_run("Gerado em 16/06/2026 às 17:00 BRT · Fonte: Medware + Kommo")
    obs_run.font.size = Pt(8)
    obs_run.italic = True
    obs_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ----- Resumo (tabela 2 colunas) -----
    doc.add_paragraph()  # espaço
    tbl_resumo = doc.add_table(rows=1 + len(RESUMO), cols=2)
    tbl_resumo.autofit = False
    tbl_resumo.columns[0].width = Cm(7)
    tbl_resumo.columns[1].width = Cm(3)

    # cabeçalho
    hdr = tbl_resumo.rows[0]
    hdr.cells[0].text = "Indicador"
    hdr.cells[1].text = "Valor"
    for c in hdr.cells:
        shade_cell(c, "1F4E79")
        set_cell_font(c, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for i, (k, v) in enumerate(RESUMO, start=1):
        row = tbl_resumo.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_font(row.cells[0], size_pt=10)
        set_cell_font(row.cells[1], size_pt=10)
        if i % 2 == 0:
            for c in row.cells:
                shade_cell(c, "F2F6FA")

    # ----- Seção Agendamentos -----
    doc.add_paragraph()
    sec = doc.add_paragraph()
    sec_run = sec.add_run("Agendamentos (ordenado por horário)")
    sec_run.bold = True
    sec_run.font.size = Pt(12)
    sec_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ----- Tabela principal -----
    cab = ["Hora", "Paciente", "Idade", "Plano", "Medware", "Etapa Kommo", "URL Kommo"]
    tbl = doc.add_table(rows=1 + len(LINHAS), cols=len(cab))
    tbl.autofit = False
    widths_cm = [1.4, 6.5, 1.3, 3.0, 2.8, 3.0, 8.4]
    for i, w in enumerate(widths_cm):
        tbl.columns[i].width = Cm(w)

    # cabeçalho
    hdr = tbl.rows[0]
    for i, txt in enumerate(cab):
        hdr.cells[i].text = txt
        shade_cell(hdr.cells[i], "1F4E79")
        set_cell_font(hdr.cells[i], size_pt=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # linhas
    for i, linha in enumerate(LINHAS, start=1):
        hora, paciente, idade, plano, medware, kommo, lead_id = linha
        row = tbl.rows[i]
        row.cells[0].text = hora
        row.cells[1].text = paciente
        row.cells[2].text = idade
        row.cells[3].text = plano
        row.cells[4].text = medware
        row.cells[5].text = kommo

        # URL clicável
        p = row.cells[6].paragraphs[0]
        if lead_id:
            url = f"https://univeja.kommo.com/leads/detail/{lead_id}"
            add_hyperlink(p, lead_id, url, color="0563C1", underline=True)
        else:
            p.add_run("—").font.size = Pt(8)

        # Formatação
        for j, c in enumerate(row.cells[:-1]):
            set_cell_font(c, size_pt=8.5)
        # Centraliza hora e idade
        for j in [0, 2]:
            for p in row.cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Cores condicionais
        if "Realizado" in medware and "5-Agendado" in kommo:
            for c in row.cells:
                shade_cell(c, "FFF8DC")  # amarelo claro
        elif "Sem Kommo" in kommo:
            for c in row.cells:
                shade_cell(c, "FFE8E8")  # rosa claro
        elif "Agendado" in medware and "Realizado" not in medware:
            for c in row.cells:
                shade_cell(c, "FFF3DC")  # laranja claro
        else:
            if i % 2 == 0:
                for c in row.cells:
                    shade_cell(c, "F7F9FB")

    # Legenda
    leg = doc.add_paragraph()
    leg_run = leg.add_run(
        "⚭ Lead conjunto (família 2+ pacientes no mesmo lead Kommo) · "
        "⚠ Atenção · "
        "Amarelo = discrepância Medware/Kommo · "
        "Rosa = sem lead Kommo · "
        "Laranja = agendado sem confirmação"
    )
    leg_run.font.size = Pt(8)
    leg_run.italic = True
    leg_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ----- Alertas críticos -----
    doc.add_paragraph()
    sec = doc.add_paragraph()
    sec_run = sec.add_run("Alertas críticos")
    sec_run.bold = True
    sec_run.font.size = Pt(12)
    sec_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for titulo, corpo in ALERTAS:
        p = doc.add_paragraph()
        p_run = p.add_run(titulo)
        p_run.bold = True
        p_run.font.size = Pt(9.5)
        p_run.font.color.rgb = RGBColor(0xA8, 0x2A, 0x1F)

        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(0.5)
        cp_run = cp.add_run(corpo)
        cp_run.font.size = Pt(9)
        cp_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Rodapé / nota final
    doc.add_paragraph()
    rod = doc.add_paragraph()
    rod_run = rod.add_run(
        "Próxima ação sugerida: sincronizar etapas Kommo dos 2 leads em discrepância "
        "(Ryan/Iara + Mariana) com status_id 91486864 (8-REALIZADO CONSULTA). "
        "Cadastrar lead Kommo para Milza de Castro Santana. "
        "Conferir procedimento de Bento Caetano (5 meses)."
    )
    rod_run.font.size = Pt(8)
    rod_run.italic = True
    rod_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(OUTPUT)
    print(f"Word gerado: {OUTPUT}")


if __name__ == "__main__":
    gerar()
