#!/usr/bin/env python3
"""Gera documento Word: Alinhamento de Agrupadores de Procedimentos ↔
Campos do Kommo ↔ Códigos do Medware.

Documento de referência única para secretária, Lia e desenvolvedores
saberem qual agrupador disparar em cada situação clínica.
"""
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE


# ============================================================
# OS 4 AGRUPADORES — fonte única de verdade
# ============================================================

AGRUPADORES = [
    {
        "n": 1,
        "nome": "Adulto / Rotina",
        "subtitulo": "Paciente ≥ 3 anos · Consulta de rotina",
        "kommo_perfil": [
            "Criança de 3 a 12 anos",
            "Adolescente de 13 a 18 anos",
            "Adulto de 19 a 49",
            "Acima de 50 anos",
        ],
        "kommo_motivo_palavras": [
            "rotina", "anual", "renovar receita", "primeira consulta",
            "check-up", "preventiva",
        ],
        "exames": [
            ("41301250", "Mapeamento de Retina de Ambos os Olhos"),
            ("41501128", "Paquimetria Ultrassônica de Ambos os Olhos"),
            ("41501012", "Biometria Ultrassônica de Ambos os Olhos"),
            ("41301200", "Exame de Motilidade"),
            ("41301323", "Tonometria"),
            ("41301315", "Retinografia Monocular Ambos os Olhos"),
            ("41301080", "Ceratoscopia Computadorizada de Ambos os Olhos"),
            ("41401271", "Teste de Cores de Ambos os Olhos"),
            ("41301269", "Microscopia Especular de Córnea"),
        ],
    },
    {
        "n": 2,
        "nome": "Adulto / Emergência",
        "subtitulo": "Paciente ≥ 3 anos · Consulta de emergência",
        "kommo_perfil": [
            "Criança de 3 a 12 anos",
            "Adolescente de 13 a 18 anos",
            "Adulto de 19 a 49",
            "Acima de 50 anos",
        ],
        "kommo_motivo_palavras": [
            "urgente", "emergência", "dor forte", "trauma",
            "corpo estranho", "sangramento", "perdi a visão",
        ],
        "exames": [
            ("41501128", "Paquimetria Ultrassônica de Ambos os Olhos"),
            ("41501012", "Biometria Ultrassônica de Ambos os Olhos"),
            ("41301200", "Exame de Motilidade"),
            ("41301323", "Tonometria"),
            ("41301080", "Ceratoscopia Computadorizada de Ambos os Olhos"),
            ("41301269", "Microscopia Especular de Córnea"),
        ],
    },
    {
        "n": 3,
        "nome": "Criança / Rotina",
        "subtitulo": "Paciente < 3 anos · Consulta de rotina",
        "kommo_perfil": ["Bebê de 0 a 2 anos"],
        "kommo_motivo_palavras": [
            "rotina", "primeira consulta", "preventiva", "check-up",
        ],
        "exames": [
            ("41301250", "Mapeamento de Retina de Ambos os Olhos"),
            ("41501128", "Paquimetria Ultrassônica de Ambos os Olhos"),
            ("41501012", "Biometria Ultrassônica de Ambos os Olhos"),
            ("41301200", "Exame de Motilidade"),
            ("41301323", "Tonometria"),
            ("41301315", "Retinografia Monocular Ambos os Olhos"),
        ],
    },
    {
        "n": 4,
        "nome": "Criança / Urgência",
        "subtitulo": "Paciente < 3 anos · Consulta de urgência",
        "kommo_perfil": ["Bebê de 0 a 2 anos"],
        "kommo_motivo_palavras": [
            "urgente", "olho vermelho", "trauma", "sangramento",
            "muito choro", "secreção",
        ],
        "exames": [
            ("41501128", "Paquimetria Ultrassônica de Ambos os Olhos"),
            ("41501012", "Biometria Ultrassônica de Ambos os Olhos"),
            ("41301200", "Exame de Motilidade"),
            ("41301323", "Tonometria"),
            ("41301315", "Retinografia Monocular Ambos os Olhos"),
        ],
    },
]


KOMMO_CAMPOS = [
    ("1.PERFIL 1º PACIENTE", "1257961", "multiselect",
     "Faixa etária do paciente principal — eixo 1 do agrupador. "
     "5 enums: Bebê 0-2, Criança 3-12, Adolescente 13-18, "
     "Adulto 19-49, Acima de 50."),
    ("1.MOTIVO CONSULTA", "1255727", "textarea",
     "Texto livre descrito pelo paciente. Lia/Haiku detecta palavras "
     "de urgência pra decidir agrupador (eixo 2)."),
    ("ESPECIALID", "1259130", "multiselect",
     "Especialidade do atendimento (Catarata, Retina, etc). Não é "
     "o que decide o agrupador, mas direciona médico."),
    ("MÉDICOS", "1256257", "multiselect",
     "Dr. Karla, Dr. Fabricio, etc. Define codMedico no Medware."),
    ("UNIDADE", "1245125", "select",
     "Asa Norte ou Águas Claras. Define codUnidade no Medware."),
    ("AÇÕES/CORRIGIR", "1259312", "multiselect",
     "Tem opção 'Urgente' — sinal humano de que agrupador é 2 ou 4 "
     "mesmo se o motivo livre não detonou palavra de urgência."),
]


REGRA_DE_NEGOCIO = """
1. SECRETÁRIA NÃO ESCOLHE EXAMES MANUALMENTE
   Em vez de marcar exame por exame na hora de agendar, ela seleciona
   UM agrupador. O sistema vincula automaticamente os exames do
   protocolo.

2. EIXO 1 — FAIXA ETÁRIA (corte em 3 anos)
   < 3 anos → Agrupador 3 ou 4
   ≥ 3 anos → Agrupador 1 ou 2
   Fonte preferencial: data de nascimento (calcula idade real).
   Fonte fallback: campo "1.PERFIL 1º PACIENTE" do Kommo.

3. EIXO 2 — MOTIVO DA CONSULTA (corte em urgência)
   Rotina → Agrupador 1 ou 3 (protocolo completo de check-up)
   Urgência/Emergência → Agrupador 2 ou 4 (foco em diagnóstico
   imediato, pula exames eletivos)
   Fonte: campo "1.MOTIVO CONSULTA" + flag "Urgente" em AÇÕES.

4. SELEÇÃO AUTOMÁTICA PELA LIA
   Quando o agente conversa com o paciente e identifica idade +
   motivo, chama `selecionar_agrupador(perfil_kommo, birth_date_iso,
   motivo)` no Python e usa o resultado pra montar o body do POST
   Medware com codProcedimentos = lista de cods.

5. AUDITORIA E AJUSTE MANUAL
   Em casos atípicos, atendente humana pode trocar o agrupador no
   Kommo via campo dedicado (a criar: "AGRUPADOR ESCOLHIDO") e a
   gravação refaz com a lista nova.
"""


# ============================================================
# GERAÇÃO DO WORD
# ============================================================

def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def add_p(doc, text, *, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    return p


def add_capa(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BLINK OFTALMOLOGIA")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agrupadores de Procedimentos")
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Alinhamento Kommo ↔ Medware ↔ Lia")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Reduz tempo operacional da secretaria · "
        "Evita esquecimento de exames obrigatórios · "
        "Lia escolhe agrupador automaticamente baseado em "
        "idade + motivo da consulta"
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento gerado em " + date.today().strftime("%d/%m/%Y"))
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()


def add_regra_negocio(doc):
    add_h1(doc, "Regra de negócio")
    for parag in REGRA_DE_NEGOCIO.strip().split("\n\n"):
        add_p(doc, parag.strip())
    doc.add_paragraph()
    doc.add_page_break()


def add_tabela_resumo(doc):
    add_h1(doc, "Resumo dos 4 agrupadores")
    add_p(
        doc,
        "Visão consolidada da matriz idade × motivo:",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    doc.add_paragraph()

    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Light Grid Accent 1"
    headers = tbl.rows[0].cells
    headers[0].text = ""
    headers[1].text = "ROTINA"
    headers[2].text = "EMERGÊNCIA / URGÊNCIA"
    for c in headers:
        for r in c.paragraphs[0].runs:
            r.bold = True

    row1 = tbl.rows[1].cells
    row1[0].text = "≥ 3 ANOS\n(adulto)"
    row1[1].text = "Agrupador 1\n9 exames (protocolo completo)"
    row1[2].text = "Agrupador 2\n6 exames focais"

    row2 = tbl.rows[2].cells
    row2[0].text = "< 3 ANOS\n(bebê)"
    row2[1].text = "Agrupador 3\n6 exames (sem ceratoscopia/cores/microscopia)"
    row2[2].text = "Agrupador 4\n5 exames (sem mapa retina)"

    for c in tbl.rows[1].cells + tbl.rows[2].cells:
        for r in c.paragraphs[0].runs:
            r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_page_break()


def add_agrupador_detalhado(doc, ag):
    add_h2(doc, f"Agrupador {ag['n']} — {ag['nome']}")
    add_p(doc, ag["subtitulo"], italic=True, size=11,
          color=RGBColor(0x77, 0x77, 0x77))

    # Kommo perfil
    r = doc.add_paragraph().add_run("Campo Kommo 1.PERFIL 1º PACIENTE: ")
    r.bold = True
    r.font.size = Pt(11)
    doc.paragraphs[-1].add_run(
        ", ".join(ag["kommo_perfil"])
    ).font.size = Pt(11)

    # Kommo motivo
    r = doc.add_paragraph().add_run(
        "Palavras-chave que disparam (1.MOTIVO CONSULTA): "
    )
    r.bold = True
    r.font.size = Pt(11)
    doc.paragraphs[-1].add_run(
        ", ".join(f'"{p}"' for p in ag["kommo_motivo_palavras"])
    ).font.size = Pt(11)

    doc.add_paragraph()

    # Tabela de exames
    r = doc.add_paragraph().add_run(
        f"Exames vinculados ({len(ag['exames'])} no total):"
    )
    r.bold = True
    r.font.size = Pt(11)

    tbl = doc.add_table(rows=len(ag["exames"]) + 1, cols=2)
    tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells
    h[0].text = "Cód. Medware"
    h[1].text = "Nome do exame"
    for c in h:
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)

    for i, (cod, nome) in enumerate(ag["exames"], start=1):
        row = tbl.rows[i].cells
        row[0].text = cod
        row[1].text = nome
        for c in row:
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_page_break()


def add_campos_kommo(doc):
    add_h1(doc, "Campos relevantes no Kommo")
    add_p(
        doc,
        "Quais custom_fields do CRM são lidos pela Lia ou pela "
        "secretaria pra decidir o agrupador.",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    doc.add_paragraph()

    tbl = doc.add_table(rows=len(KOMMO_CAMPOS) + 1, cols=4)
    tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells
    h[0].text = "Campo"
    h[1].text = "field_id"
    h[2].text = "Tipo"
    h[3].text = "Uso na seleção do agrupador"
    for c in h:
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)

    for i, (nome, fid, tipo, uso) in enumerate(KOMMO_CAMPOS, start=1):
        row = tbl.rows[i].cells
        row[0].text = nome
        row[1].text = fid
        row[2].text = tipo
        row[3].text = uso
        for c in row:
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_page_break()


def add_proximos_passos(doc):
    add_h1(doc, "Próximos passos para integrar no agente")
    passos = [
        "Importar `voice_agent.procedimentos.selecionar_agrupador` "
        "em `voice_agent.agendamento.executar_agendamento`.",

        "Após resolver cod_medico/cod_unidade, calcular agrupador "
        "passando perfil_kommo (de caller_context['known']) + "
        "birth_date_iso (do paciente) + motivo (de 1.MOTIVO CONSULTA).",

        "Adaptar `voice_agent.medware.criar_agendamento` para aceitar "
        "lista de codProcedimentos em vez de um só. Iterar e fazer "
        "POST de cada exame (ou ajustar para o endpoint correto "
        "de múltiplos exames no Medware se existir).",

        "Criar novo campo Kommo opcional 'AGRUPADOR ESCOLHIDO' "
        "(select com 4 opções) — permite atendente humano sobrescrever "
        "a escolha automática em casos atípicos.",

        "Adicionar `/admin/dry-agrupador?perfil=X&birth=Y&motivo=Z` "
        "que devolve o agrupador escolhido + lista de cods. Útil "
        "pra a secretaria validar a regra antes de gravar.",

        "Atualizar `_MASTER_INSTRUCTION.md` da Lia: instruir a "
        "perguntar idade do paciente E motivo da consulta antes de "
        "confirmar agendamento, e mencionar quais exames vão ser "
        "realizados (transparência).",

        "Adicionar enum 'Urgência detectada' em '1.MOTIVO CONSULTA' "
        "do Kommo se preferir validação humana antes da seleção "
        "automática (alternativa mais conservadora).",
    ]
    for i, p in enumerate(passos, start=1):
        para = doc.add_paragraph()
        r = para.add_run(f"  {i:02d}.  ")
        r.font.name = "Courier New"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x99, 0x33)
        para.add_run(p).font.size = Pt(11)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_capa(doc)
    add_regra_negocio(doc)
    add_tabela_resumo(doc)
    add_h1(doc, "Detalhamento dos agrupadores")
    for ag in AGRUPADORES:
        add_agrupador_detalhado(doc, ag)
    add_campos_kommo(doc)
    add_proximos_passos(doc)

    out = (
        "/sessions/dazzling-bold-davinci/mnt/outputs/"
        "alinhamento_agrupadores_procedimentos.docx"
    )
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
