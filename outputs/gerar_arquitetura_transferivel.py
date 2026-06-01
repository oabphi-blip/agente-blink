#!/usr/bin/env python3
"""Gera documento Word: Arquitetura defensiva da Lia → reutilizável
no agente de autorização do portal dos convênios.

Sintetiza tudo o que aprendemos hoje (30/05/2026) em 7 players
arquiteturais transferíveis, anti-padrões a evitar, e checklist de
bootstrap pra novo agente.
"""
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# CONTEÚDO
# ============================================================

CAPA = {
    "titulo": "Arquitetura Defensiva de Agente IA",
    "subtitulo": (
        "Padrões transferíveis da Lia (atendimento WhatsApp) "
        "para o próximo agente: autorização de convênios"
    ),
    "autoria": "Blink Oftalmologia · Equipe Fábio Philipe",
    "data": date.today().strftime("%d/%m/%Y"),
    "epigrafe": (
        "Cada bug que custou 4 horas hoje vira teste de 30 segundos "
        "amanhã. Cada erro silencioso vira alerta antes de o cliente "
        "perceber. Cada campo deletado no painel externo vira "
        "blacklist auto-aprendida em runtime."
    ),
}

INTRO = (
    "Esse documento empacota as habilidades arquiteturais que blindaram "
    "a Lia (assistente WhatsApp da Blink) em produção real, para serem "
    "reutilizadas no próximo agente — o de autorização de convênios "
    "via portal. O objetivo é evitar repetir os mesmos 5–6 erros de "
    "diagnóstico e perseguição de causa que custaram horas em incidentes."
    "\n\n"
    "A premissa central: um agente de IA com efeito em mundo real "
    "(que conversa com paciente, grava no CRM, agenda no PEP, "
    "submete pedido em portal) precisa de SETE camadas defensivas. "
    "Sem elas, qualquer mudança externa (campo deletado no Kommo, "
    "enum alterado no Medware, layout do portal mudou) vira incidente "
    "silencioso que ninguém percebe até o cliente reclamar."
)


PLAYERS = [
    {
        "n": 1,
        "nome": "Memória ativa documentada",
        "subtitulo": "Conhecimento que sobrevive entre sessões",
        "o_que_faz": (
            "Arquivos markdown estruturados em local conhecido que descrevem "
            "cada bug histórico, sua causa raiz, a regra criada pra não "
            "repetir, e onde está o pytest que blinda. O próximo Claude/dev "
            "lê automaticamente no boot da sessão e começa contextualizado."
        ),
        "como_implementar": (
            "Pasta `agente-autorizacao/memoria/bugs-licoes/` com 1 arquivo "
            ".md por incidente. Estrutura padrão: Causa raiz · Sintoma "
            "observado · Diagnóstico falso perseguido · Métodos que "
            "evitariam · Pytest que blinda · Regra de ouro pós-incidente. "
            "Esse documento que você está lendo é exemplo. Atualizar a cada "
            "incidente, sem fail."
        ),
        "beneficio": (
            "Mata o ciclo \"toda hora esqueço o que aconteceu\". O próximo "
            "agente lê os 8–10 arquivos da Lia antes de errar igual."
        ),
        "estado_blink": "Implementado (lia-atendimento-blink/memoria/bugs-licoes/)",
    },
    {
        "n": 2,
        "nome": "Ambiente de teste isolado",
        "subtitulo": "Reproduzir bug sem cliente real",
        "o_que_faz": (
            "Endpoints administrativos que disparam o pipeline completo do "
            "agente com input controlado — sem precisar paciente real "
            "mandar mensagem, sem precisar enviar autorização real ao "
            "convênio. Devolvem o que o agente FARIA, sem efeito "
            "colateral, ou rodam o efeito real em sandbox."
        ),
        "como_implementar": (
            "Pra Lia criamos cinco endpoints: /admin/simulate-inbound "
            "(pipeline real ou dry-run), /admin/debug-extract (mostra o "
            "que Haiku extrai), /admin/dry-sync (payload Kommo sem postar), "
            "/admin/force-resync (sync real num lead arbitrário), "
            "/admin/schema-check (lista campos órfãos blacklistados). "
            "Pro agente de autorização: análogos seriam /admin/simulate-"
            "pedido, /admin/preview-formulario, /admin/dry-submit, "
            "/admin/force-resubmit, /admin/schema-portal."
        ),
        "beneficio": (
            "Diagnóstico que tomava 4 horas (caçar nos logs Easypanel + "
            "pedir paciente mandar oi + interpretar) passa a tomar 30 "
            "segundos. Esse foi o multiplicador absurdo de produtividade "
            "que descobrimos hoje."
        ),
        "estado_blink": "Implementado e validado em produção",
    },
    {
        "n": 3,
        "nome": "Pytest com cenários históricos",
        "subtitulo": "Cada bug vira teste que nunca pode regredir",
        "o_que_faz": (
            "Toda vez que descobrimos uma causa nova de falha, escrevemos "
            "1 pytest que reproduz aquele cenário exato. Antes de qualquer "
            "deploy, pytest roda — se algum cenário histórico falha, deploy "
            "é bloqueado. Garante que nunca voltamos pro mesmo bug."
        ),
        "como_implementar": (
            "Pasta `tests/` com arquivos nomeados por área: "
            "`test_filtros_lia.py`, `test_whatsapp_inbound_pipeline.py`, "
            "`test_kommo_auto_skip.py`. Cada classe TestX cobre um "
            "incidente. Pra agente de autorização: "
            "`test_formulario_portal.py`, `test_retry_envio_autorizacao.py`, "
            "`test_campos_obrigatorios_convenio.py`. "
            "Executar com `python -m pytest tests/ -v` antes de cada commit."
        ),
        "beneficio": (
            "Trava regressão. Sem pytest, mesma equipe acaba reintroduzindo "
            "bug que já foi corrigido (vimos isso no commit task #20 que "
            "removeu uma função sem remover a chamada e quebrou Lia 5 horas)."
        ),
        "estado_blink": "12 cenários cobrindo bugs históricos",
    },
    {
        "n": 4,
        "nome": "Self-healing: auto-skip com retry inteligente",
        "subtitulo": "Agente aprende com a falha externa em runtime",
        "o_que_faz": (
            "Quando o sistema externo (Kommo, portal do convênio) rejeita "
            "uma operação com erro específico (\"campo inválido\", \"enum "
            "não suportado\", \"layout mudou\"), o agente detecta no "
            "response, identifica QUAL parte do payload causou, marca como "
            "morta numa blacklist em memória, e RETENTA sem aquela parte. "
            "Próximas chamadas pulam direto."
        ),
        "como_implementar": (
            "Padrão: `_DEAD_FIELDS: set[str] = set()` class-level. Em cada "
            "submit, primeiro filtrar payload removendo campos em "
            "_DEAD_FIELDS. Se PATCH retornar 400 com error pointing pra "
            "campo X, adicionar X em _DEAD_FIELDS e retry. Até N retries. "
            "Pra portal: se autorização retornar \"campo CID obrigatório\" "
            "ou \"plano não cadastrado\", aprender e retry com fallback."
        ),
        "beneficio": (
            "Mata classe inteira \"campo foi deletado/renomeado no sistema "
            "externo e o agente continuou tentando\". Sem self-healing, "
            "qualquer mudança externa = incidente até deploy de fix. Com "
            "self-healing = degradação graciosa automática."
        ),
        "estado_blink": "Implementado (Kommo auto-skip, 5 pytest)",
    },
    {
        "n": 5,
        "nome": "Observabilidade com painel semafórico",
        "subtitulo": "Equipe vê estado sem ler logs",
        "o_que_faz": (
            "Endpoint /admin/healthz devolve JSON consolidado com "
            "timestamps da última atividade bem-sucedida em cada camada "
            "(última msg recebida, último PATCH ok, último envio ao "
            "portal, último erro 4xx, blacklist atual). Equipe abre URL "
            "e vê \"Lia processou há 30s, último PATCH 200 há 1min, 2 "
            "campos órfãos\" — sem precisar tail de logs."
        ),
        "como_implementar": (
            "Marca timestamps em Redis em cada passo importante: "
            "`SETEX blink:healthz:last_kommo_patch_ok TTL=7d <ts>`. "
            "Endpoint lê e devolve. Pra agente de autorização: "
            "marcar último envio sucedido, último 4xx, último captcha, "
            "blacklist de campos rejeitados, contadores 24h."
        ),
        "beneficio": (
            "Equipe (Stephany, Ariany, recepção) consegue ver se o "
            "agente está vivo SEM precisar Fábio. Bug avisa a si mesmo "
            "antes de virar incêndio."
        ),
        "estado_blink": "Implementado (/admin/healthz)",
    },
    {
        "n": 6,
        "nome": "Subagentes para análise paralela",
        "subtitulo": "Mapear código em 30s em vez de chutar 4h",
        "o_que_faz": (
            "Em vez de o Claude principal ficar lendo arquivo por arquivo "
            "e chutando hipótese, ele dispara subagentes Plan/Explore que "
            "leem código em paralelo e devolvem o MAPA em <200 palavras. "
            "Resposta precisa em <30s. Multiplica capacidade analítica."
        ),
        "como_implementar": (
            "Disponível na ferramenta Task do Claude. Ao receber bug, "
            "antes de qualquer ação, sempre spawnar Plan agent com: "
            "\"mapeie o caminho completo de X, liste todos pontos de "
            "descarte silencioso, identifique condições file:line\". "
            "Pra agente de autorização: mapear caminho do submit ao "
            "portal, todos os pontos onde request pode falhar silencioso."
        ),
        "beneficio": (
            "Hoje usamos isso 3 vezes — em cada incidente. Cada uso "
            "economizou 1–2 horas de chute. Substitui hipótese por "
            "análise estática."
        ),
        "estado_blink": "Usado liberalmente (Plan agent + Explore agent)",
    },
    {
        "n": 7,
        "nome": "Schema source-of-truth + boot fail-loud",
        "subtitulo": "Container não sobe se estado externo divergiu",
        "o_que_faz": (
            "No boot do agente, fazer GET em endpoint do sistema externo "
            "que lista os campos/enums atuais. Bater com o que está "
            "hardcoded no código. Se diverge, ou ABORTAR o boot com erro "
            "claro, ou marcar os divergentes como mortos (combinar com "
            "player #4). Sem isso, qualquer deletion no sistema externo "
            "vira PATCH 400 silencioso."
        ),
        "como_implementar": (
            "Pra Kommo: GET /api/v4/leads/custom_fields no startup, "
            "comparar enum_ids. Pra portal de convênio: GET no formulário "
            "padrão na primeira chamada do dia, comparar campos. "
            "Pra Medware: GET /listar_planos no startup, comparar "
            "codPlanos. Se divergir, log fail-loud + alertar Slack + "
            "blacklistar campos órfãos."
        ),
        "beneficio": (
            "Detecta divergência ANTES de a primeira chamada real falhar. "
            "Equipe vê alerta \"3 enums divergentes no Kommo\" em vez de "
            "\"PATCH falhou pra 28 pacientes hoje\"."
        ),
        "estado_blink": "Próximo passo (parcial via auto-skip do player #4)",
    },
]


ANTI_PADROES = [
    {
        "titulo": "Adivinhar path em vez de checar",
        "descricao": (
            "\"Acho que o arquivo de skill fica em ~/.claude/skills/.\" "
            "Errado. Sempre rodar `find` ou `ls` antes de copiar arquivo. "
            "Tempo: 5s. Custou 30min de chute no passado."
        ),
    },
    {
        "titulo": "Codificar mapeamento sem listar a fonte",
        "descricao": (
            "Hardcodar 7 enums quando o sistema tem 27. Lia silenciou "
            "20 convênios por meses. Antes de hardcodar lookup, listar "
            "o catálogo oficial via API/MCP."
        ),
    },
    {
        "titulo": "Múltiplas mudanças sem smoke test entre cada",
        "descricao": (
            "Editar pipeline + agendamento + responder + KB em sequência "
            "sem testar Medware no meio. Só descobre erro com paciente "
            "real. Após cada arquivo tocado, validar isoladamente."
        ),
    },
    {
        "titulo": "Mudar prompt sem rodar pytest",
        "descricao": (
            "Editar _MASTER_INSTRUCTION sem validar que regras antigas "
            "continuam disparando. Após qualquer edit em prompt/KB, "
            "rodar `python -m pytest tests/ -v` antes de commit."
        ),
    },
    {
        "titulo": "Commitar segredos",
        "descricao": (
            "CPF de paciente em commit. Token GitHub em script shell. "
            "Antes de cada commit, varrer diff por regex CPF "
            "(`\\d{11}`) e token (`ghp_[A-Za-z0-9]{36}`)."
        ),
    },
    {
        "titulo": "Confiar no deploy verde sem validar runtime",
        "descricao": (
            "Easypanel marca \"Aplicativo implantado\" como sucesso, mas "
            "o build em si usou cache antigo. Container roda código de "
            "ontem. Sempre validar com 1 chamada que prova runtime novo "
            "(endpoint novo, log novo, header com git_sha)."
        ),
    },
    {
        "titulo": "Persegir causa em vez de mapear todos pontos de falha",
        "descricao": (
            "Hoje perdi 5h perseguindo Meta webhook quando a causa real "
            "era NameError em responder.py. Antes de perseguir hipótese, "
            "spawnar Plan agent pra mapear TODOS pontos de descarte "
            "silencioso. Diagnóstico vira eliminação, não chute."
        ),
    },
]


CHECKLIST_BOOTSTRAP = [
    "Criar pasta do projeto + CLAUDE.md com regras críticas do domínio",
    "Setup git + remote GitHub + token em Keychain (NUNCA em commit)",
    "Stack mínima: FastAPI + httpx + pytest + redis-py + python-docx",
    "Dockerfile com cache-bust ARG pra forçar rebuild quando preciso",
    "Easypanel: app + redis + auto-deploy GitHub webhook",
    "Pasta `tests/` desde o dia 1 — escrever 1 teste pro happy path",
    "Pasta `memoria/bugs-licoes/` com README explicando estrutura",
    "Endpoint `/health` mínimo: status + configs + integrações",
    "Endpoint `/admin/healthz` com timestamps das últimas operações",
    "Endpoint `/admin/simulate-X` que dispara pipeline com input controlado",
    "Endpoint `/admin/dry-X` que mostra o que SERIA enviado sem efeito",
    "Endpoint `/admin/force-X` que executa real num registro arbitrário",
    "Endpoint `/admin/schema-check` que lista campos blacklistados",
    "Padrão auto-skip + retry no client de cada sistema externo",
    "Marcar timestamps Redis em cada operação relevante",
    "Plan agent disparado a cada bug ANTES de qualquer hipótese",
    "Atualizar CLAUDE.md anti-padrões a cada incidente",
    "Atualizar memoria/bugs-licoes a cada incidente",
    "Smoke-test em CI (GitHub Action) chamando /admin/simulate antes de deploy verde",
    "Alerta Slack a cada 4xx/timeout/exception no client externo",
    "Schema-on-startup que valida hardcoded vs sistema externo real",
]


REGRA_DE_OURO = (
    "Antes de adicionar mais um log de debug em produção, antes de "
    "perseguir mais uma hipótese, antes de fazer mais um deploy: "
    "spawnar Plan agent pra mapear o caminho inteiro do dado. Listar "
    "todos pontos de descarte silencioso. Perguntar \"qual o teste "
    "real que reproduz o sintoma sem cliente real?\". Se a resposta "
    "não está em <30s, está em <2h depois de implementar o ambiente "
    "de teste — não em <8h debugando logs no escuro."
)


# ============================================================
# GERAÇÃO DO WORD
# ============================================================

def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def add_p(doc, text, *, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    return p


def add_capa(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BLINK OFTALMOLOGIA")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CAPA["titulo"])
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CAPA["subtitulo"])
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("“" + CAPA["epigrafe"] + "”")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CAPA["autoria"])
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento gerado em " + CAPA["data"])
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()


def add_intro(doc):
    add_h1(doc, "Por que esse documento existe")
    for parag in INTRO.split("\n\n"):
        add_p(doc, parag.strip())
    doc.add_paragraph()
    doc.add_page_break()


def add_players(doc):
    add_h1(doc, "Os 7 players da arquitetura defensiva")
    add_p(
        doc,
        "Cada player corresponde a uma camada de proteção. Implementar "
        "os 7 mata 95% dos incidentes silenciosos. Implementar 3 "
        "primeiros já reduz drasticamente o tempo de diagnóstico.",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    doc.add_paragraph()

    for p in PLAYERS:
        add_h2(doc, f"Player #{p['n']} — {p['nome']}")
        add_p(doc, p["subtitulo"], italic=True, size=11,
              color=RGBColor(0x77, 0x77, 0x77))

        r = doc.add_paragraph().add_run("O que faz: ")
        r.bold = True
        r.font.size = Pt(11)
        doc.paragraphs[-1].add_run(p["o_que_faz"]).font.size = Pt(11)

        r = doc.add_paragraph().add_run("Como implementar: ")
        r.bold = True
        r.font.size = Pt(11)
        doc.paragraphs[-1].add_run(p["como_implementar"]).font.size = Pt(11)

        r = doc.add_paragraph().add_run("Benefício: ")
        r.bold = True
        r.font.size = Pt(11)
        doc.paragraphs[-1].add_run(p["beneficio"]).font.size = Pt(11)

        r = doc.add_paragraph().add_run("Estado na Lia (Blink): ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x99, 0x33)
        doc.paragraphs[-1].add_run(p["estado_blink"]).font.size = Pt(10)

        doc.add_paragraph()

    doc.add_page_break()


def add_anti_padroes(doc):
    add_h1(doc, "Anti-padrões a evitar")
    add_p(
        doc,
        "Os 7 erros mais frequentes que descobrimos hoje e que "
        "documentamos no CLAUDE.md da Lia. Cada um custou pelo menos "
        "1 hora de retrabalho.",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    doc.add_paragraph()

    for i, ap in enumerate(ANTI_PADROES, start=1):
        add_h2(doc, f"{i}. {ap['titulo']}")
        add_p(doc, ap["descricao"])
        doc.add_paragraph()

    doc.add_page_break()


def add_checklist(doc):
    add_h1(doc, "Checklist de bootstrap")
    add_p(
        doc,
        "21 itens pra implementar nos primeiros dias do agente de "
        "autorização. Implementar nessa ordem garante que cada player "
        "está presente antes do agente atender em produção.",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    doc.add_paragraph()

    for i, item in enumerate(CHECKLIST_BOOTSTRAP, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"  [{i:02d}]  ")
        r.font.name = "Courier New"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x99, 0x33)
        p.add_run(item).font.size = Pt(11)

    doc.add_page_break()


def add_regra_de_ouro(doc):
    add_h1(doc, "A regra de ouro")
    p = doc.add_paragraph()
    r = p.add_run("“" + REGRA_DE_OURO + "”")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_capa(doc)
    add_intro(doc)
    add_players(doc)
    add_anti_padroes(doc)
    add_checklist(doc)
    add_regra_de_ouro(doc)

    out = (
        "/sessions/dazzling-bold-davinci/mnt/outputs/"
        "arquitetura_defensiva_agente_blink.docx"
    )
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
